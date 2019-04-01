from breakdown.models import Ticket, Breakdown, Status, FunctionGroup, BreakdownCategory
from breakdown import serializers as bds
from rest_framework import viewsets
from django.http import Http404
from rest_framework.decorators import action, api_view
from rest_framework.response import Response
from rest_framework.reverse import reverse
from django_filters.rest_framework import DjangoFilterBackend
from rest_framework.views import APIView
from rest_framework_jwt.views import obtain_jwt_token
from django.contrib.auth.models import User
from .filters import TicketFilter
from jira import JIRA
from django.conf import settings
from docx import Document
from docx.shared import Inches
from openpyxl import Workbook
from openpyxl.styles import Fill, Border, Side, PatternFill, Alignment, Font, Color
from openpyxl.styles.colors import WHITE
from django.db.models import Q
from docx.shared import RGBColor
from datetime import datetime
from common.exception_handler import DataUnavailable
from common.config import SysConfig, ConfigKey
import os

DOC_PATH = settings.MEDIA_ROOT + 'breakdown/'
sysConfig = SysConfig()

@api_view(['GET'])
def api_root(request, format=None):
    return Response({
        'tickets': reverse('ticket-list', request=request, format=format),
        'breakdowns': reverse('breakdown-list', request=request, format=format),
    })

class TicketViewSet(viewsets.ModelViewSet):
    queryset = Ticket.objects.all()
    serializer_class = bds.TicketSerializer
    filter_backends = (DjangoFilterBackend,)
    filter_class = TicketFilter

    def perform_create(self, serializer):        
        ticket = serializer.save()
        categoryIds = [int(s) for s in sysConfig.getStringConfig(ConfigKey.BK_TEMPLATE_CATEGORY_IDS).split(',')]
        inFds = sysConfig.getStringConfig(ConfigKey.BK_TEMPLATE_INFD_VALUES).split(',')
        inBks = sysConfig.getStringConfig(ConfigKey.BK_TEMPLATE_INBK_VALUES).split(',')
        staus = int(sysConfig.getStringConfig(ConfigKey.BK_TEMPLATE_STATUS))
        func = int(sysConfig.getStringConfig(ConfigKey.BK_TEMPLATE_FUNC))
        for i in range(len(categoryIds)):
            self.getBreakdown(
                i+1, BreakdownCategory.objects.get(pk=categoryIds[i]), 
                ticket,
                Status.objects.get(pk=staus),
                FunctionGroup.objects.get(pk=func),
                True if inFds[i] == '1' else False, 
                True if inBks[i] == '1' else False, 
                ).save()

    def getBreakdown(self, sequence, category, ticket, status, func, infd, inbk):
        bk = Breakdown()
        bk.sequence = sequence
        bk.ticket = ticket
        bk.category = category
        bk.status = status
        bk.function_group = func
        bk.in_fd = infd
        bk.in_bk = inbk
        return bk



class BreakdownViewSet(viewsets.ModelViewSet):
    
    queryset = Breakdown.objects.all()
    serializer_class = bds.BreakdownSerializer

    @action(detail=True)
    def breakdowns(self, request, pk=None):
        serializer = self.get_serializer(self.queryset.filter(ticket=pk), many=True)
        return Response(serializer.data)

class UploadImageView(APIView):
    queryset = Breakdown.objects.none()
    def get_object(self, pk):
        try:
            return Breakdown.objects.get(pk=pk)
        except Breakdown.DoesNotExist:
            raise Http404

    def put(self, request, pk, format=None):    
        bk = self.get_object(pk)    
        if (bk.image1 == '' or bk.image1 == None):          
            bk.image1 = self.request.FILES.get('file')
        elif (bk.image2 == '' or bk.image2 == None):
            bk.image2 = self.request.FILES.get('file')
        elif (bk.image3 == '' or bk.image3 == None):
            bk.image3 = self.request.FILES.get('file')
        bk.save()
        bks = bds.BreakdownSerializer(bk)
        return Response(bks.data)
    
    def delete(self, request, pk, format=None):
        filename = self.request.query_params.get('name')
        bk = self.get_object(pk)
        if (bk.image1 != '' and bk.image1 != None and bk.image1.url.endswith(filename)):
            os.remove(settings.MEDIA_ROOT + bk.image1.url)
            bk.image1 = ''
        elif (bk.image2 != '' and bk.image2 != None and bk.image2.url.endswith(filename)):
            os.remove(settings.MEDIA_ROOT + bk.image2.url)
            bk.image2 = ''
        elif (bk.image3 != '' and bk.image3 != None and bk.image3.url.endswith(filename)):
            os.remove(settings.MEDIA_ROOT + bk.image3.url)
            bk.image3 = ''

        bk.save()
        return Response('{success:true}')


class StatusViewSet(viewsets.ReadOnlyModelViewSet):
    queryset = Status.objects.all()
    serializer_class = bds.StatusSerializer
    pagination_class = None

class FuncGroupsViewSet(viewsets.ReadOnlyModelViewSet):
    queryset = FunctionGroup.objects.all()
    serializer_class = bds.FunctionGroupSerializer
    pagination_class = None

class UsersViewSet(viewsets.ReadOnlyModelViewSet):
    queryset = User.objects.filter(~Q(id = 1)).order_by('username')
    serializer_class = bds.UserSerializer
    pagination_class = None

class CategoryViewSet(viewsets.ReadOnlyModelViewSet):
    # queryset = BreakdownCategory.objects.filter(Q(parent=None))
    queryset = BreakdownCategory.objects.all()
    serializer_class = bds.CategorySerializer
    pagination_class = None


class JiraView(APIView):

    queryset = Breakdown.objects.none()
    serializer_class = None
    def get(self, request, format=None):
        """
        Return a list of all issues.
        """
        jira_server = sysConfig.getStringConfig(ConfigKey.JIRA_SERVER)
        jira_username = sysConfig.getStringConfig(ConfigKey.JIRA_ACCOUNT)
        jira_password = sysConfig.getStringConfig(ConfigKey.JIRA_PASSWORD)

        myjira = JIRA(jira_server,basic_auth=(jira_username,jira_password))
        # jql_str = "project = IPROS AND createdDate >= 2018-01-01 AND type = Task"
        jql_str = sysConfig.getStringConfig(ConfigKey.JIRA_QUERY)
        fields = [
            'status',
            'fixVersions',
            'assignee',
            'key',
            'summary',
            'duedate',
            'customfield_10112',
            'id'
        ]
        data = myjira.search_issues(jql_str, 0, -1, True, fields, None, True)
        return Response(data)

class UploadToJiraView(APIView):

    queryset = Breakdown.objects.none()
    serializer_class = None
    def get_ticket(self, pk):
        try:
            return Ticket.objects.get(pk=pk)
        except Ticket.DoesNotExist:
            raise Http404

    def get(self, request, pk, format=None):
        """
        Return a list of all issues.
        """
        jira_server = sysConfig.getStringConfig(ConfigKey.JIRA_SERVER)
        jira_username = sysConfig.getStringConfig(ConfigKey.JIRA_ACCOUNT)
        jira_password = sysConfig.getStringConfig(ConfigKey.JIRA_PASSWORD)        

        myjira = JIRA(jira_server,basic_auth=(jira_username,jira_password))
        ticket = self.get_ticket(pk)
        docPath = DOC_PATH + ticket.ticket_no + '/'
        docName = sysConfig.getStringConfig(ConfigKey.DOC_FD_NAME).replace('{0}', ticket.ticket_no)
        fullDocName = docPath + docName + '.docx'
        xlsName = sysConfig.getStringConfig(ConfigKey.DOC_BK_NAME).replace('{0}', ticket.ticket_no)
        fullXlsName = docPath + xlsName + '.xlsx'
        now = datetime.now()
        date_time = now.strftime("%Y%m%d%H%M%S")
        try:
            if os.path.exists(fullDocName):
                myjira.add_attachment(ticket.jira_id, fullDocName, docName + '-' + date_time + '.docx')
            if os.path.exists(fullXlsName):
                myjira.add_attachment(ticket.jira_id, fullXlsName, xlsName + '-' + date_time + '.xlsx')
            return Response({'success'})
        except:
            return Response({'failed'})
        

class DocumentListView(APIView):
    queryset = Breakdown.objects.none()
    serializer_class = None

    def get(self, request, format=None):
        """
        Return a list of all docs.
        """
        data = []
        ticketFolders = os.listdir(DOC_PATH)
        for folder in ticketFolders:
            fdPath = DOC_PATH + folder + '/'
            docName = sysConfig.getStringConfig(ConfigKey.DOC_FD_NAME).replace('{0}', folder) + '.docx'
            docs = []
            if os.path.exists(fdPath + docName):
                docs.append(docName)                
            
            xlsName = sysConfig.getStringConfig(ConfigKey.DOC_BK_NAME).replace('{0}', folder) + '.xlsx'
            if os.path.exists(fdPath + xlsName):
                docs.append(xlsName)  

            data.append({
                    'ticket': folder,
                    'docs':docs
                })
        return Response(data)

class GenerateDocView(APIView):
    queryset = Breakdown.objects.none()
    def get_ticket(self, pk):
        try:
            return Ticket.objects.get(pk=pk)
        except Ticket.DoesNotExist:
            raise Http404
    
    def get_breakdown(self, pk):
        try:
            return Breakdown.objects.filter(ticket=pk)
        except Breakdown.DoesNotExist:
            raise Http404

    def get(self, request, pk, format=None):
        ticket = self.get_ticket(pk)
        breakdowns = self.get_breakdown(pk)
        if breakdowns == None or len(breakdowns) == 0:
            raise DataUnavailable('No breakdown found!', 'Data_Unavailable')

        docPath = DOC_PATH + ticket.ticket_no + '/'
        docNameStr = sysConfig.getStringConfig(ConfigKey.DOC_FD_NAME).replace('{0}', ticket.ticket_no)
        docName = docPath + docNameStr + '.docx'
        xlsNameStr = sysConfig.getStringConfig(ConfigKey.DOC_BK_NAME).replace('{0}', ticket.ticket_no)
        xlsName = docPath + xlsNameStr + '.xlsx'
        self.generateDoc(ticket, breakdowns.filter(in_fd=1), docName, docPath)
        self.generateXls(ticket, breakdowns.filter(in_bk=1), xlsName, docPath)
        return Response(docName)

    def generateDoc(self, ticket, breakdowns, docName, docPath):
        print('generating doc..........')
        document = Document()
        document.add_heading('FD - ' + ticket.ticket_no, 0)
        document.add_paragraph(ticket.summary)
        currentSubCategory = None
        currentFuncGroup = None
        currentParentCat = None
        styles = document.styles
        font1 = styles['Heading 1'].font
        font1.color.rgb = RGBColor(0xEA, 0x6F, 0x5A)

        font4 = styles['Heading 4'].font
        font4.color.rgb = RGBColor(0x00, 0x00, 0x00)
        for bk in breakdowns:
            category = bk.category
            funcGroup = bk.function_group
            if category == None or funcGroup == None:
                raise DataUnavailable('Category or Function Group is not defined!', 'Data_Unavailable')
            # Add parent category
            if currentParentCat == None:
                if category.parent == None:
                    currentParentCat = category.id
                    currentSubCategory = category.id
                    currentFuncGroup = None
                    document.add_heading(category.code, level=1)
                else:
                    currentParentCat = category.parent.id
                    document.add_heading(category.parent.code, level=1)
                    currentFuncGroup = None
            else:
                if category.parent == None:
                    currentParentCat = category.id
                    currentSubCategory = category.id
                    currentFuncGroup = None
                    document.add_heading(category.code, level=1)
                elif currentParentCat != category.parent.id:
                    currentParentCat = category.parent.id
                    document.add_heading(category.parent.code, level=1)
                    currentFuncGroup = None
            # Add sub category
            if currentSubCategory == None or currentSubCategory != category.id:
                currentSubCategory = category.id
                currentFuncGroup = None
                document.add_heading(category.code, level=2)
            # Add function group
            if funcGroup != None and funcGroup.description != '---' and (currentFuncGroup == None or currentFuncGroup != funcGroup.id):
                currentFuncGroup = funcGroup.id
                document.add_heading(funcGroup.description, level=4)
            try:
                startIndex = bk.description.index('{')
                endIndex = bk.description.index('}') + 1
            except:
                startIndex = 0
                endIndex = 0
            if startIndex > 0 or endIndex > 0:
                document.add_paragraph(bk.description[endIndex:], style='List Bullet')
            else:
                document.add_paragraph(bk.description, style='List Bullet')
            if bk.image1 != '':
                document.add_picture(bk.image1, width=Inches(5.0))

            if bk.image2 != '':
                document.add_picture(bk.image2, width=Inches(5.0))

            if bk.image3 != '':
                document.add_picture(bk.image3, width=Inches(5.0))
        if os.path.exists(docName):
            os.remove(docName)
        if not os.path.isdir(docPath):
            os.mkdir(docPath)
        document.save(docName)
    
    def generateXls(self, ticket, breakdowns, docName, docPath):
        print('generating excel..........')
        wb = Workbook()
        ws = wb.active
        ws.title = 'Effort Breakdown'
        titleFont = Font(name='Arial', size=12, bold=True, color=WHITE)
        titleFill = PatternFill(patternType='solid', fill_type = 'solid', fgColor=Color('409EFF'))
        thin_border = Border(left=Side(style='thin'), 
                     right=Side(style='thin'), 
                     top=Side(style='thin'), 
                     bottom=Side(style='thin'))

        detailFont = Font(name='Arial', size=10, bold=False)
        startCol = 2
        startRow = 2
        ws.column_dimensions['A'].width = 2
        ws.column_dimensions['B'].width = 5
        ws.column_dimensions['C'].width = 15
        # ws.column_dimensions['D'].width = 25
        ws.column_dimensions['D'].width = 80
        ws.column_dimensions['E'].width = 15
        self.createCell(ws, startRow, startCol, titleFont, titleFill, thin_border, 'SN')
        self.createCell(ws, startRow, startCol+1, titleFont, titleFill, thin_border, 'Category')
        # self.createCell(ws, startRow, startCol+2, titleFont, titleFill, thin_border, 'Function Group')
        self.createCell(ws, startRow, startCol+2, titleFont, titleFill, thin_border, 'Items')
        self.createCell(ws, startRow, startCol+3, titleFont, titleFill, thin_border, 'Effort(mds)')

        currentSubCategory = None
        currentFuncGroup = None
        categoryStartRow = 0
        funcStartRow = 0
        i = 0
        devEffort = 0
        row = startRow
        for bk in breakdowns:
            i = i + 1
            row = row + 1
            self.createCell(ws, row, startCol, detailFont, None, thin_border, i)
            category = bk.category
            funcGroup = bk.function_group
            if category == None or funcGroup == None:
                raise DataUnavailable('Category or Function Group is not defined!', 'Data_Unavailable')
            if category.parent != None:
                category = category.parent
            self.createCell(ws, row, startCol+1, detailFont, None, thin_border, category.code)
            # self.createCell(ws, row, startCol+2, detailFont, None, thin_border, funcGroup.description)
            try:
                startIndex = bk.description.index('{') + 1
                endIndex = bk.description.index('}')
            except:
                startIndex = 0
                endIndex = 0
            if startIndex > 0 or endIndex > 0:
                self.createCell(ws, row, startCol+2, detailFont, None, thin_border, bk.description[startIndex:endIndex])
            else:
                self.createCell(ws, row, startCol+2, detailFont, None, thin_border, bk.description)

            self.createCell(ws, row, startCol+3, detailFont, None, thin_border, bk.effort)
            devEffort += bk.effort
            
            # if currentFuncGroup == None:
            #     funcStartRow = row
            #     currentFuncGroup = funcGroup.id
            # elif currentFuncGroup != funcGroup.id:
            #     if funcStartRow < row - 1:
            #         ws.merge_cells(start_row=funcStartRow, start_column=startCol+2, end_row=row - 1, end_column=startCol+2)
            #     funcStartRow = row
            #     currentFuncGroup = funcGroup.id

            if currentSubCategory == None:
                categoryStartRow = row
                currentSubCategory = category.id
            elif currentSubCategory != category.id:
                if categoryStartRow < row - 1:
                    ws.merge_cells(start_row=categoryStartRow, start_column=startCol+1, end_row=row - 1, end_column=startCol+1)
                categoryStartRow = row
                currentSubCategory = category.id
                # set index for func group
                # if funcStartRow < row - 1:
                #     ws.merge_cells(start_row=funcStartRow, start_column=startCol+2, end_row=row - 1, end_column=startCol+2)
                # funcStartRow = row
                currentFuncGroup = funcGroup.id
            elif i == len(breakdowns) + 1 and categoryStartRow < row:
                ws.merge_cells(start_row=categoryStartRow, start_column=startCol+1, end_row=row, end_column=startCol+1)
                # set index for func group
                # if funcStartRow < row:
                #     ws.merge_cells(start_row=funcStartRow, start_column=startCol+2, end_row=row, end_column=startCol+2)

        formula = '=SUM(F' + str(startRow) + ':F' + str(row) + ')'
        self.createCell(ws, row + 1, startCol+3, titleFont, titleFill, None, formula)

        if os.path.exists(docName):
            os.remove(docName)
        if not os.path.isdir(docPath):
            os.mkdir(docPath)

        wb.save(docName)

    def createCell(self, ws, row, col, font=None, fill=None, border=None, value=''):
        col_sn = ws.cell(row=row, column=col, value=value)
        if font:
            col_sn.font = font
        if fill:
            col_sn.fill = fill
        if border:
            col_sn.border = border
        col_sn.alignment = Alignment(wrap_text=True)





